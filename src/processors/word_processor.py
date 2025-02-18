from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from .base.document_processor import DocumentProcessor
import tempfile
import os
import docx
import textract

class WordProcessor(DocumentProcessor):
    """Word document (.doc, .docx) processor implementation"""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=int(os.getenv('CHUNK_SIZE', '1000')),
            chunk_overlap=int(os.getenv('CHUNK_OVERLAP', '200'))
        )

    def process(self, file_obj):
        # Handle both string paths and file-like objects
        if isinstance(file_obj, str):
            file_path = file_obj
        else:
            # Create a temporary file for the uploaded content
            with tempfile.NamedTemporaryFile(delete=False, suffix=self._get_suffix(file_obj)) as tmp_file:
                if hasattr(file_obj, 'name'):
                    file_path = file_obj.name
                else:
                    content = file_obj.read() if hasattr(file_obj, 'read') else file_obj
                    if isinstance(content, str):
                        content = content.encode('utf-8')
                    tmp_file.write(content)
                    file_path = tmp_file.name

        try:
            # Extract text based on file type
            if file_path.lower().endswith('.docx'):
                text = self._process_docx(file_path)
            else:  # .doc
                text = self._process_doc(file_path)

            # Create a single document with the extracted text
            doc = Document(page_content=text, metadata={"source": file_path})

            # Split into chunks
            chunks = self.text_splitter.split_documents([doc])

            return chunks
        finally:
            # Clean up temporary file if created
            if 'tmp_file' in locals():
                os.unlink(file_path)

    def _process_docx(self, file_path):
        """Process .docx file using python-docx"""
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    def _process_doc(self, file_path):
        """Process .doc file using textract"""
        return textract.process(file_path).decode('utf-8')

    def _get_suffix(self, file_obj):
        """Determine file suffix based on filename"""
        if hasattr(file_obj, 'name'):
            return os.path.splitext(file_obj.name)[1].lower()
        return '.docx'  # default to .docx if unable to determine